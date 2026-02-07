"""
Word document (.docx) text extraction using python-docx.

Extracts:
- Paragraphs
- Tables (converted to readable text)
- Headers and footers
"""

import io
import logging

from docx import Document
from docx.table import Table

logger = logging.getLogger(__name__)


def extract_docx_text(source: bytes | str) -> str:
    """
    Extract text content from a Word document (.docx).

    Args:
        source: Raw DOCX file content (bytes) or path to DOCX file (str)

    Returns:
        Extracted plain text with paragraphs separated by newlines.
        Returns empty string if extraction fails.
    """
    try:
        # Support both bytes and file paths
        file_input = io.BytesIO(source) if isinstance(source, bytes) else source
        doc = Document(file_input)
        text_parts: list[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # Extract tables
        for table in doc.tables:
            table_text = _table_to_text(table)
            if table_text:
                text_parts.append(table_text)

        # Extract headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text and text not in text_parts:
                        text_parts.insert(0, f"[Header] {text}")

            # Footer
            if section.footer:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text and text not in text_parts:
                        text_parts.append(f"[Footer] {text}")

        full_text = "\n\n".join(text_parts)
        logger.info("Extracted %d characters from DOCX", len(full_text))
        return full_text

    except Exception as e:
        logger.error("Failed to extract DOCX: %s", e)
        return ""


def _table_to_text(table: Table) -> str:
    """
    Convert a Word table to readable text format.

    Uses first row as headers if it looks like a header row.
    """
    rows = table.rows
    if not rows:
        return ""

    lines: list[str] = []

    # Extract all cells
    all_rows = []
    for row in rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):  # Skip empty rows
            all_rows.append(cells)

    if not all_rows:
        return ""

    # Use first row as headers
    headers = all_rows[0]

    # Process data rows
    for i, row in enumerate(all_rows[1:], start=1):
        pairs = []
        for j, cell in enumerate(row):
            if cell:
                header = headers[j] if j < len(headers) else f"Col{j+1}"
                pairs.append(f"{header}={cell}")

        if pairs:
            lines.append(f"Row {i}: {', '.join(pairs)}")

    return "\n".join(lines)
